"""
简历解析服务
"""
import os
import uuid
from typing import Dict, Any, Optional
from datetime import datetime
import docx
from pydantic import BaseModel, Field

from app.services.llm_factory import llm_factory
from app.models.database import LLMConfig
from app.core.config import settings


class ResumeData(BaseModel):
    """简历数据结构"""
    name: str = Field(default="", description="姓名")
    phone: Optional[str] = Field(default=None, description="电话")
    email: Optional[str] = Field(default=None, description="邮箱")
    location: Optional[str] = Field(default=None, description="所在地")
    education: list = Field(default_factory=list, description="教育背景")
    experience: list = Field(default_factory=list, description="工作经历")
    skills: list = Field(default_factory=list, description="技能")
    projects: list = Field(default_factory=list, description="项目经历")
    summary: Optional[str] = Field(default=None, description="个人总结")


class ResumeParser:
    """简历解析器"""
    
    def __init__(self, llm_config: Optional[LLMConfig] = None):
        self.llm_config = llm_config
    
    def parse_docx(self, file_path: str) -> ResumeData:
        """
        解析Word简历
        
        Args:
            file_path: 简历文件路径
            
        Returns:
            ResumeData: 解析后的简历数据
        """
        # 1. 提取原始文本
        doc = docx.Document(file_path)
        raw_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        # 2. 提取表格内容
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    tables_text.append(" | ".join(row_text))
        
        full_text = raw_text + "\n" + "\n".join(tables_text)
        
        # 3. 使用LLM结构化提取
        parsed_data = self._extract_with_llm(full_text)
        
        return ResumeData(**parsed_data)
    
    def _extract_with_llm(self, text: str) -> Dict[str, Any]:
        """使用LLM提取结构化信息"""
        from langchain_openai import ChatOpenAI
        import json
        import re
        
        # kimik2.5 支持 128K 上下文，可以处理完整简历
        # 只在极端情况下截断（超过 100K 字符）
        max_length = 100000
        if len(text) > max_length:
            truncated = text[:max_length]
            last_period = max(truncated.rfind('。'), truncated.rfind('\n'), truncated.rfind('.'))
            if last_period > max_length * 0.9:
                text = truncated[:last_period + 1]
            else:
                text = truncated
            print(f"简历文本过长，已截断至 {len(text)} 字符")
        else:
            print(f"简历文本长度: {len(text)} 字符，使用完整文本解析")
        
        prompt = f"""请从以下简历文本中提取结构化信息，以JSON格式返回。

简历文本：
{text}

请提取以下字段（如果某项不存在，返回空字符串或空列表）：
{{
    "name": "姓名",
    "phone": "电话",
    "email": "邮箱",
    "location": "所在地",
    "education": [
        {{
            "school": "学校名称",
            "major": "专业",
            "degree": "学位",
            "start_date": "开始时间",
            "end_date": "结束时间"
        }}
    ],
    "experience": [
        {{
            "company": "公司名称",
            "position": "职位",
            "department": "部门",
            "start_date": "开始时间",
            "end_date": "结束时间",
            "description": "工作描述"
        }}
    ],
    "skills": ["技能1", "技能2"],
    "projects": [
        {{
            "name": "项目名称",
            "description": "项目描述",
            "technologies": ["技术1", "技术2"],
            "role": "个人角色",
            "achievements": ["成果1", "成果2"]
        }}
    ],
    "summary": "个人总结/自我评价"
}}

重要提示：
1. 只返回JSON格式数据，不要其他说明文字
2. 日期格式统一为：YYYY-MM
3. 确保JSON格式正确，可以被解析
4. 所有字符串必须使用双引号，不能使用单引号
5. 确保所有括号正确闭合
"""
        
        try:
            # 创建LLM实例
            if self.llm_config:
                # 使用传入的配置
                llm = llm_factory.create_llm(self.llm_config)
            else:
                # 使用默认配置（从settings读取）
                api_key = settings.DEFAULT_LLM_API_KEY
                base_url = settings.DEFAULT_LLM_BASE_URL
                model = settings.DEFAULT_LLM_MODEL
                
                if not api_key:
                    print("未配置LLM API Key，使用规则提取")
                    return self._extract_with_rules(text)
                
                # 检测是否为 Kimi/Moonshot 模型
                model_lower = model.lower()
                is_kimi = "kimi" in model_lower or "moonshot" in model_lower
                
                # 构建参数
                llm_params = {
                    "model": model,
                    "api_key": api_key,
                    "base_url": base_url,
                }
                
                # 对于 Kimi/Moonshot 模型，必须显式设置 temperature=1（使用整数）
                if is_kimi:
                    llm_params["temperature"] = 1  # 使用整数 1 而不是 1.0
                    print(f"[DEBUG] Kimi/Moonshot model in resume_parser, setting temperature=1")
                else:
                    llm_params["temperature"] = 0.7
                
                print(f"[DEBUG] Resume parser creating ChatOpenAI with params: {llm_params}")
                llm = ChatOpenAI(**llm_params)
            
            response = llm.invoke(prompt)
            content = response.content if hasattr(response, 'content') else str(response)
            
            # 解析JSON - 改进提取逻辑
            data = self._parse_json_response(content)
            if data:
                return data
            else:
                print("JSON解析失败，使用规则提取")
                return self._extract_with_rules(text)
            
        except Exception as e:
            print(f"LLM解析失败: {e}，使用规则提取")
            return self._extract_with_rules(text)
    
    def _parse_json_response(self, content: str) -> Optional[Dict[str, Any]]:
        """解析LLM返回的JSON响应"""
        import json
        
        # 尝试多种方式提取JSON
        # 方法1: 提取 ```json 代码块
        if '```json' in content:
            try:
                json_str = content.split('```json')[1].split('```')[0].strip()
                return json.loads(json_str)
            except:
                pass
        
        # 方法2: 提取 ``` 代码块
        if '```' in content:
            try:
                json_str = content.split('```')[1].split('```')[0].strip()
                return json.loads(json_str)
            except:
                pass
        
        # 方法3: 查找JSON对象（从第一个 { 到最后一个 }）
        try:
            start = content.find('{')
            end = content.rfind('}')
            if start != -1 and end != -1 and end > start:
                json_str = content[start:end+1]
                return json.loads(json_str)
        except:
            pass
        
        # 方法4: 直接解析整个内容
        try:
            return json.loads(content.strip())
        except:
            pass
        
        return None
    
    def _extract_with_rules(self, text: str) -> Dict[str, Any]:
        """使用规则提取（fallback方案）"""
        import re
        
        data = {
            "name": "",
            "phone": "",
            "email": "",
            "location": "",
            "education": [],
            "experience": [],
            "skills": [],
            "projects": [],
            "summary": ""
        }
        
        # 提取姓名（简单规则：第一行或包含"姓名"）
        lines = text.split('\n')
        if lines:
            data["name"] = lines[0].strip()[:20]  # 取第一行作为姓名
        
        # 提取电话
        phone_pattern = r'1[3-9]\d{9}'
        phone_match = re.search(phone_pattern, text)
        if phone_match:
            data["phone"] = phone_match.group()
        
        # 提取邮箱
        email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
        email_match = re.search(email_pattern, text)
        if email_match:
            data["email"] = email_match.group()
        
        return data


# 全局解析器实例
resume_parser = ResumeParser()
